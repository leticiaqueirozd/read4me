import os
from flask import Flask, render_template, request
from transformers import BlipProcessor, BlipForConditionalGeneration
from gtts import gTTS
from werkzeug.utils import secure_filename
from PIL import Image
import pdfplumber
import pytesseract
import docx
from googletrans import Translator

app = Flask(__name__)

processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
translator = Translator()

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'jpg', 'jpeg', 'png', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def describe_image(image_path, language):
    image = Image.open(image_path)
    inputs = processor(image, return_tensors="pt")
    outputs = model.generate(**inputs)
    description = processor.decode(outputs[0], skip_special_tokens=True)
    
    if language != 'en':
        translated_description = translate_text(description, language)
        return translated_description
    return description

def translate_text(text, dest_language):
    try:
        translated = translator.translate(text, dest=dest_language)
        return translated.text
    except Exception as e:
        print(f"Erro na tradução: {e}")
        return text

def extract_text_from_pdf(file_path):
    full_text = ""
    image_descriptions = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += page_text + "\n"
            else:
                img_data = page.to_image()
                img_path = "temp_image.png"
                img_data.save(img_path)
                full_text += "\n[Texto extraído por OCR]: " + pytesseract.image_to_string(img_path)

            for img in page.images:
                img_path = extract_image_from_pdf(page, img)
                description = describe_image(img_path)
                if description:
                    image_descriptions.append(description)
    
    return full_text, image_descriptions

def extract_text_from_docx(file_path):
    full_text = ""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

def extract_image_from_pdf(page, img):
    x0, y0, x1, y1 = img['x0'], img['y0'], img['x1'], img['y1']
    img_data = page.to_image()
    pil_image = img_data.original.convert("RGB")
    cropped_img = pil_image.crop((x0, y0, x1, y1))
    img_path = 'extracted_image.png'
    cropped_img.save(img_path)
    return img_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_file():
    file = request.files['file']
    language = request.form.get('language', 'pt')
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            full_text = ""
            image_descriptions = []

            # Processando PDF
            if file.filename.endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            full_text += page_text + "\n"
                        else:
                            img_data = page.to_image()
                            img_path = "temp_image.png"
                            img_data.save(img_path)
                            full_text += "\n[Texto extraído por OCR]: " + pytesseract.image_to_string(img_path)
                        for img in page.images:
                            img_path = extract_image_from_pdf(page, img)
                            description = describe_image(img_path, language)
                            image_descriptions.append(description)

            elif file.filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as txt_file:
                    full_text = txt_file.read()

            elif file.filename.endswith('.docx'):
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    full_text += para.text + "\n"

            elif file.filename.endswith(('jpg', 'jpeg', 'png')):
                image = Image.open(file)
                text = pytesseract.image_to_string(image)
                description = describe_image(file_path, language)
                full_text = text + description

            tts = gTTS(full_text, lang=language)
            audio_path = 'static/audio/output.mp3'
            tts.save(audio_path)

            full_text += "\n".join(image_descriptions)

            return render_template('result.html', text=full_text, audio_file='audio/output.mp3', image_descriptions=image_descriptions)

        finally:
            os.remove(file_path)

    return "Arquivo inválido ou tipo de arquivo não permitido."

if __name__ == '__main__':
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    if not os.path.exists("static/audio"):
        os.makedirs("static/audio")
    app.run(debug=True)
